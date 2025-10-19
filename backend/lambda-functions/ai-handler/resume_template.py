"""
This module creates a resume template as a Word document.
It will be used as a template for the AI-generated resume content.
"""

import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

def create_resume_template():
    """Create a resume template as a Word document and return it as bytes."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add name placeholder
    name = doc.add_paragraph()
    name_run = name.add_run("[FULL_NAME]")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add contact info placeholder
    contact = doc.add_paragraph()
    contact_run = contact.add_run("[CONTACT_INFO]")
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a line separator
    doc.add_paragraph("_" * 80)
    
    # Add Professional Summary section
    summary_heading = doc.add_heading('Professional Summary', level=1)
    summary_heading.style.font.size = Pt(14)
    summary_heading.style.font.bold = True
    doc.add_paragraph("[PROFESSIONAL_SUMMARY]")
    
    # Add Skills section
    skills_heading = doc.add_heading('Skills', level=1)
    skills_heading.style.font.size = Pt(14)
    skills_heading.style.font.bold = True
    doc.add_paragraph("[SKILLS]")
    
    # Add Experience section
    exp_heading = doc.add_heading('Experience', level=1)
    exp_heading.style.font.size = Pt(14)
    exp_heading.style.font.bold = True
    doc.add_paragraph("[EXPERIENCE]")
    
    # Add Education section
    edu_heading = doc.add_heading('Education', level=1)
    edu_heading.style.font.size = Pt(14)
    edu_heading.style.font.bold = True
    doc.add_paragraph("[EDUCATION]")
    
    # Save to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

if __name__ == "__main__":
    # For testing: save the template to a file
    with open("resume_template.docx", "wb") as f:
        f.write(create_resume_template())
    print("Template created and saved as resume_template.docx")